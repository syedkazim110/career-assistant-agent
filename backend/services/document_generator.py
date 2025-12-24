from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
import os
import logging
import re
from typing import Optional, List, Tuple
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Service for generating resume and cover letter documents"""
    
    def __init__(self, output_dir: str = "generated"):
        """
        Initialize document generator
        
        Args:
            output_dir: Directory to save generated documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def parse_markdown_text(self, text: str) -> List[Tuple[str, bool, bool]]:
        """
        Parse markdown text and return list of (text, is_bold, is_italic) tuples
        
        Args:
            text: Text with markdown formatting
            
        Returns:
            List of tuples with text and formatting flags
        """
        result = []
        i = 0
        
        while i < len(text):
            # Check for bold (**text**)
            if i < len(text) - 1 and text[i:i+2] == '**':
                # Find the closing **
                end = text.find('**', i + 2)
                if end != -1:
                    bold_text = text[i+2:end]
                    result.append((bold_text, True, False))
                    i = end + 2
                    continue
            
            # Check for italic (*text* or _text_)
            if text[i] in ['*', '_']:
                marker = text[i]
                end = text.find(marker, i + 1)
                if end != -1 and end - i > 1:
                    italic_text = text[i+1:end]
                    result.append((italic_text, False, True))
                    i = end + 1
                    continue
            
            # Regular text - find next markdown marker
            next_marker = len(text)
            for marker_pos in [text.find('**', i), text.find('*', i), text.find('_', i)]:
                if marker_pos != -1 and marker_pos < next_marker:
                    next_marker = marker_pos
            
            if next_marker > i:
                regular_text = text[i:next_marker]
                if regular_text:
                    result.append((regular_text, False, False))
                i = next_marker
            else:
                # No more markers, add rest of text
                remaining = text[i:]
                if remaining:
                    result.append((remaining, False, False))
                break
        
        return result
    
    def generate_docx_resume(self, content: str, filename: str) -> str:
        """
        Generate a resume in DOCX format
        
        Args:
            content: Resume content
            filename: Name for the output file
            
        Returns:
            Path to the generated file
        """
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Split content into sections
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if it's a heading (starts with ## or ###)
                if line.startswith('##'):
                    # Add heading
                    heading_text = line.replace('#', '').strip()
                    heading = doc.add_heading(heading_text, level=1)
                    heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
                    heading.runs[0].font.size = Pt(14)
                    heading.runs[0].font.bold = True
                elif line.startswith('- ') or line.startswith('• '):
                    # Bullet point - parse markdown within bullet
                    bullet_text = line[2:].strip()
                    para = doc.add_paragraph(style='List Bullet')
                    para.paragraph_format.left_indent = Inches(0.25)
                    
                    # Parse markdown formatting in bullet text
                    parts = self.parse_markdown_text(bullet_text)
                    for text, is_bold, is_italic in parts:
                        run = para.add_run(text)
                        run.font.bold = is_bold
                        run.font.italic = is_italic
                else:
                    # Normal paragraph - parse markdown formatting
                    para = doc.add_paragraph()
                    para.paragraph_format.space_after = Pt(6)
                    
                    # Parse markdown formatting
                    parts = self.parse_markdown_text(line)
                    for text, is_bold, is_italic in parts:
                        run = para.add_run(text)
                        if is_bold:
                            run.font.bold = True
                        if is_italic:
                            run.font.italic = True
            
            # Save document
            file_path = os.path.join(self.output_dir, filename)
            doc.save(file_path)
            
            logger.info(f"Resume DOCX generated: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX resume: {str(e)}")
            raise
    
    def convert_markdown_to_html(self, text: str) -> str:
        """
        Convert markdown formatting to HTML for ReportLab
        
        Args:
            text: Text with markdown formatting
            
        Returns:
            HTML formatted text
        """
        # First escape special HTML characters
        text = text.replace('&', '&amp;')
        
        # Convert **bold** to <b>bold</b>
        while '**' in text:
            first = text.find('**')
            if first == -1:
                break
            second = text.find('**', first + 2)
            if second == -1:
                break
            text = text[:first] + '<b>' + text[first+2:second] + '</b>' + text[second+2:]
        
        # Convert *italic* to <i>italic</i>
        parts = []
        i = 0
        while i < len(text):
            if text[i] == '*' and (i == 0 or text[i-1] != '<'):
                end = text.find('*', i + 1)
                if end != -1:
                    parts.append('<i>' + text[i+1:end] + '</i>')
                    i = end + 1
                    continue
            parts.append(text[i])
            i += 1
        
        return ''.join(parts) if parts else text
    
    def generate_pdf_resume(self, content: str, filename: str) -> str:
        """
        Generate a resume in PDF format
        
        Args:
            content: Resume content
            filename: Name for the output file
            
        Returns:
            Path to the generated file
        """
        try:
            file_path = os.path.join(self.output_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(
                file_path,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='CustomHeading',
                parent=styles['Heading1'],
                fontSize=14,
                textColor=RGBColor(0, 51, 102),
                spaceAfter=12,
                spaceBefore=12,
                bold=True
            ))
            styles.add(ParagraphStyle(
                name='CustomBody',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=6
            ))
            
            # Build document content
            story = []
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.1*inch))
                    continue
                
                # Check if it's a heading
                if line.startswith('##'):
                    heading_text = line.replace('#', '').strip()
                    # Escape and convert markdown
                    heading_text = self.convert_markdown_to_html(heading_text)
                    story.append(Paragraph(heading_text, styles['CustomHeading']))
                else:
                    # Convert markdown to HTML for ReportLab
                    formatted_line = self.convert_markdown_to_html(line)
                    story.append(Paragraph(formatted_line, styles['CustomBody']))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"Resume PDF generated: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating PDF resume: {str(e)}")
            raise
    
    def generate_docx_cover_letter(self, content: str, filename: str) -> str:
        """
        Generate a cover letter in DOCX format
        
        Args:
            content: Cover letter content
            filename: Name for the output file
            
        Returns:
            Path to the generated file
        """
        try:
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add content - skip date as AI should include it
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue
                
                # Parse markdown formatting for cover letter
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(6)
                
                # Parse markdown formatting
                parts = self.parse_markdown_text(line)
                for text, is_bold, is_italic in parts:
                    run = para.add_run(text)
                    if is_bold:
                        run.font.bold = True
                    if is_italic:
                        run.font.italic = True
            
            # Save document
            file_path = os.path.join(self.output_dir, filename)
            doc.save(file_path)
            
            logger.info(f"Cover letter DOCX generated: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating DOCX cover letter: {str(e)}")
            raise
    
    def generate_pdf_cover_letter(self, content: str, filename: str) -> str:
        """
        Generate a cover letter in PDF format
        
        Args:
            content: Cover letter content
            filename: Name for the output file
            
        Returns:
            Path to the generated file
        """
        try:
            file_path = os.path.join(self.output_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(
                file_path,
                pagesize=letter,
                rightMargin=1*inch,
                leftMargin=1*inch,
                topMargin=1*inch,
                bottomMargin=1*inch
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='CoverLetterBody',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                alignment=TA_JUSTIFY
            ))
            
            # Build document content
            story = []
            
            # Skip adding date if AI already includes it in content
            # Add content
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 0.1*inch))
                    continue
                
                # Convert markdown to HTML for ReportLab
                formatted_line = self.convert_markdown_to_html(line)
                story.append(Paragraph(formatted_line, styles['CoverLetterBody']))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"Cover letter PDF generated: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error generating PDF cover letter: {str(e)}")
            raise
